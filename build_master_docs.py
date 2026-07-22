from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "docs" / "master"
NAVY = "17365D"; BLUE = "2E74B5"; LIGHT = "EAF0F7"; GRAY = "667085"


def set_font(run, size=11, bold=False, italic=False, color="111827", mono=False):
    name = "Consolas" if mono else "Arial"
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def page_field(p):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run("AISTUDYTEC  •  "), 8.5, color=GRAY)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); p._p.append(fld)


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); pr.append(shd)


def margins(cell):
    pr = cell._tc.get_or_add_tcPr(); mar = OxmlElement("w:tcMar")
    for side, value in (("top",100),("start",120),("bottom",100),("end",120)):
        el=OxmlElement(f"w:{side}"); el.set(qn("w:w"),str(value)); el.set(qn("w:type"),"dxa"); mar.append(el)
    pr.append(mar)


def add_inline(p, text, size=11, color="111827", mono=False):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part: continue
        if part.startswith("`"):
            set_font(p.add_run(part[1:-1]), size-0.5, color="344054", mono=True)
        elif part.startswith("**"):
            set_font(p.add_run(part[2:-2]), size, bold=True, color=color)
        else:
            set_font(p.add_run(part), size, color=color, mono=mono)


def make_doc(path):
    lines = path.read_text(encoding="utf-8").splitlines()
    title = lines[0].lstrip("# ")
    doc = Document(); sec=doc.sections[0]
    sec.page_width=Inches(8.27); sec.page_height=Inches(11.69)
    sec.top_margin=Inches(0.9); sec.bottom_margin=Inches(0.8); sec.left_margin=Inches(0.9); sec.right_margin=Inches(0.8)
    page_field(sec.footer.paragraphs[0])
    normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"),"Arial"); normal._element.rPr.rFonts.set(qn("w:hAnsi"),"Arial")
    for name,size,color,before,after in (("Heading 1",16,NAVY,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",11.5,NAVY,9,4)):
        st=doc.styles[name]; st.font.name="Arial"; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
        st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True

    for _ in range(4): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("AISTUDYTEC"), 13, bold=True, color=BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(title.replace("AISTUDYTEC — ","")), 25, bold=True, color=NAVY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(12)
    set_font(p.add_run("Documento mestre • baseline versionável"), 12, italic=True, color=GRAY)
    for _ in range(8): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Fortaleza — CE\n2026"), 11, bold=True)
    doc.add_page_break()

    i=1; in_code=False
    while i < len(lines):
        line=lines[i]
        if line.startswith("```"):
            in_code=not in_code; i+=1; continue
        if in_code:
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.25); p.paragraph_format.space_after=Pt(0)
            add_inline(p,line,9,mono=True); i+=1; continue
        if not line.strip(): i+=1; continue
        if line.startswith("## "):
            doc.add_heading(line[3:],1); i+=1; continue
        if line.startswith("### "):
            doc.add_heading(line[4:],2); i+=1; continue
        if line.startswith("#### "):
            doc.add_heading(line[5:],3); i+=1; continue
        if line.startswith("| "):
            block=[]
            while i<len(lines) and lines[i].startswith("|"):
                block.append([c.strip() for c in lines[i].strip().strip("|").split("|")]); i+=1
            if len(block)>1 and all(set(x.replace(" ","")) <= set("-:") for x in block[1]): block.pop(1)
            cols=max(len(r) for r in block); t=doc.add_table(rows=1,cols=cols); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
            tr_pr=t.rows[0]._tr.get_or_add_trPr(); header_flag=OxmlElement("w:tblHeader"); header_flag.set(qn("w:val"),"true"); tr_pr.append(header_flag)
            for c in range(cols):
                cell=t.rows[0].cells[c]; shade(cell,LIGHT); margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
                add_inline(cell.paragraphs[0], block[0][c] if c<len(block[0]) else "", 8.5, NAVY)
                for run in cell.paragraphs[0].runs: run.bold=True
            for row in block[1:]:
                cells=t.add_row().cells
                for c in range(cols): margins(cells[c]); add_inline(cells[c].paragraphs[0],row[c] if c<len(row) else "",8.2)
            doc.add_paragraph(); continue
        if re.match(r"^\d+\. ",line) or line.startswith("- "):
            style="List Number" if re.match(r"^\d+\. ",line) else "List Bullet"
            text=re.sub(r"^(\d+\. |- )","",line)
            p=doc.add_paragraph(style=style); p.paragraph_format.space_after=Pt(4); add_inline(p,text,10.5); i+=1; continue
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.space_after=Pt(7); p.paragraph_format.line_spacing=1.25
        add_inline(p,line,11); i+=1

    doc.core_properties.title=title; doc.core_properties.author="Projeto AISTUDYTEC"
    out=path.with_suffix(".docx"); doc.save(out); print(out)


for md in sorted(SRC.glob("0*_*.md")): make_doc(md)
