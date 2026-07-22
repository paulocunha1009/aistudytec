from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "docs" / "TCC_AISTUDYTEC.docx"
NAVY = "17365D"
BLUE = "2E74B5"
LIGHT = "EAF0F7"
GRAY = "666666"


def font(run, size=12, bold=False, italic=False, color="000000", name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}")) or OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa"); tcMar.append(node)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    font(run, 9, color=GRAY)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def body(doc, text, citation=None):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.49)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text + (f" {citation}" if citation else ""))
    font(r, 11.5)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    font(p.add_run(text), 11)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def chapter(doc, number, title):
    doc.add_page_break()
    heading(doc, f"{number} {title.upper()}", 1)


def add_table(doc, headers, rows, widths=None, caption=None, source=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4)
        font(p.add_run(caption), 10, bold=True, color=NAVY)
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.style = "Table Grid"
    tr_pr = t.rows[0]._tr.get_or_add_trPr()
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    tr_pr.append(header_flag)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; shade(cell, LIGHT); set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        font(cell.paragraphs[0].add_run(h), 9.5, bold=True, color=NAVY)
        if widths: cell.width = Inches(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i]); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            font(cells[i].paragraphs[0].add_run(str(value)), 9.2)
            if widths: cells[i].width = Inches(widths[i])
    if source:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(8)
        font(p.add_run(f"Fonte: {source}"), 8.5, italic=True, color=GRAY)
    return t


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
sec.top_margin = Inches(1.18); sec.bottom_margin = Inches(0.79)
sec.left_margin = Inches(1.18); sec.right_margin = Inches(0.79)
sec.header_distance = Inches(0.45); sec.footer_distance = Inches(0.45)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"; normal.font.size = Pt(11.5)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
for name, size, color, before, after in (("Heading 1", 15, NAVY, 16, 10), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 11.5, NAVY, 9, 4)):
    st = styles[name]; st.font.name = "Arial"; st.font.size = Pt(size); st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st._element.rPr.rFonts.set(qn("w:ascii"), "Arial"); st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)

header = sec.header.paragraphs[0]
font(header.add_run("AISTUDYTEC | Relatório técnico-científico"), 8.5, color=GRAY)
add_page_number(sec.footer.paragraphs[0])

# Capa
for _ in range(5): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("AISTUDYTEC"), 26, bold=True, color=NAVY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("PLATAFORMA DE ESTUDO INDEPENDENTE COM INTELIGÊNCIA ARTIFICIAL, CURADORIA DOCENTE E ACOMPANHAMENTO POR HABILIDADES"), 16, bold=True, color=BLUE)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Trabalho técnico-científico para qualificação acadêmica"), 13, italic=True)
for _ in range(8): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Autoria: responsável pelo projeto AISTUDYTEC (identificação não fornecida)"), 11)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Fortaleza — CE\n2026"), 11, bold=True)

# Folha de rosto e nota de integridade
doc.add_page_break()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("AISTUDYTEC"), 18, bold=True, color=NAVY)
body(doc, "Trabalho de desenvolvimento tecnológico apresentado em versão para qualificação. Instituição, programa, linha de pesquisa, orientação e autoria nominal não foram informados e, por integridade acadêmica, não são presumidos neste documento.")
heading(doc, "Declaração de integridade dos dados", 2)
body(doc, "Este relatório não atribui eficácia pedagógica ao protótipo. Os únicos resultados próprios reportados são evidências técnicas reproduzíveis obtidas no repositório até 22 de julho de 2026: estrutura do software, contagem de componentes, execução de testes automatizados, compilação do frontend e smoke tests das APIs Gemini e YouTube com credenciais reais mantidas em segredo. Não foram inventados estudantes, docentes, turmas, entrevistas, questionários, notas, grupos de controle ou métricas de uso. Dados educacionais externos são identificados por fonte e ano.")
heading(doc, "Natureza e delimitação", 2)
body(doc, "O objeto é um artefato computacional em estágio de protótipo funcional. A contribuição discutida é de projeto: integração entre geração assistida por IA, curadoria humana, feedback formativo, domínio por habilidade e revisão espaçada. A validação educacional permanece como agenda de pesquisa e deve ser submetida a avaliação ética antes de envolver menores de idade.")

# Resumo
doc.add_page_break(); heading(doc, "RESUMO", 1)
body(doc, "Este trabalho apresenta o desenvolvimento e a avaliação técnica inicial do AISTUDYTEC, plataforma web destinada ao estudo independente e complementar de estudantes do Ensino Médio brasileiro. O objetivo foi projetar um artefato que preserve a centralidade docente e converta conteúdo gerado em aprendizagem ativa: três níveis de profundidade articulam curiosidade, objetivos, leitura guiada, investigação, vídeo com missão, desafio, reflexão, diário e quiz por habilidade. Adotou-se pesquisa aplicada inspirada em Design Science Research, com análise documental, levantamento bibliográfico, modelagem, implementação incremental e verificação técnica. A versão examinada contém backend Flask/SQLite, frontend React, 23 rotas, 12 tabelas, 13 componentes React, 40 testes backend e 15 testes frontend aprovados. Gemini e YouTube foram validados por smoke tests reais sem exposição das chaves. Esses resultados demonstram consistência técnica nos comportamentos cobertos, mas não demonstram impacto sobre aprendizagem. A contribuição é uma arquitetura pedagógico-técnica auditável, acompanhada de protocolo para avaliação futura.")
p = doc.add_paragraph(); font(p.add_run("Palavras-chave: inteligência artificial na educação; Ensino Médio; avaliação formativa; prática de recuperação; curadoria docente; software educacional."), 11, bold=True)

heading(doc, "ABSTRACT", 1)
body(doc, "This study presents the development and initial technical evaluation of AISTUDYTEC, a web platform for independent, complementary study by Brazilian upper-secondary students. The artifact combines teacher curation with a three-stage active-learning path: curiosity, objectives, guided reading, investigation, purposeful video viewing, hands-on production, reflection, a local journal, and skill-tagged quizzes. The inspected version comprises a Flask/SQLite backend, a React frontend, 23 HTTP routes, 12 database tables, 13 React components, 40 passing backend tests, and 15 passing frontend tests. Real Gemini and YouTube smoke tests were completed without exposing credentials. These findings establish technical evidence only and do not demonstrate learning effectiveness. Educational validation, privacy governance, accessibility, and controlled-pilot ethics remain prerequisites.")
p = doc.add_paragraph(); font(p.add_run("Keywords: artificial intelligence in education; secondary education; formative assessment; retrieval practice; teacher curation; educational software."), 11, bold=True)

# Sumário descritivo
doc.add_page_break(); heading(doc, "SUMÁRIO", 1)
for item in ["1 Introdução", "2 Contexto educacional e problema", "3 Referencial teórico", "4 Trabalhos e debates relacionados", "5 Metodologia", "6 Especificação do AISTUDYTEC", "7 Arquitetura e implementação", "8 Verificação e resultados técnicos", "9 Discussão", "10 Ética, privacidade e riscos", "11 Protocolo de avaliação futura", "12 Gestão Scrum e evolução MCP", "13 Limitações", "14 Conclusão", "Referências", "Apêndices"]:
    bullet(doc, item)

chapter(doc, 1, "Introdução")
body(doc, "A expansão de recursos digitais e de modelos generativos alterou as condições de produção e circulação do conhecimento escolar. Essa transformação não elimina problemas históricos: o Brasil permanece abaixo da média da OCDE em matemática, leitura e ciências no PISA 2022, e os resultados nacionais têm mostrado estabilidade de longo prazo desde 2009. No recorte brasileiro participaram 10.798 estudantes de 599 escolas, o que permite contextualizar o desafio sem confundir uma avaliação sistêmica com diagnóstico individual.", "(INEP, 2023; OECD, 2023)")
body(doc, "O debate sobre inteligência artificial na educação oscila entre promessas de personalização e preocupações sobre erro, viés, integridade acadêmica, proteção de dados e enfraquecimento da mediação humana. A UNESCO recomenda uma abordagem centrada nas pessoas, adequada à idade, com validação ética e pedagógica. A imprensa internacional de alta reputação registra a mesma ambivalência: docentes exploram feedback e explicações personalizadas, enquanto pesquisadores alertam para ciclos de entusiasmo tecnológico, desigualdade e perda de relações humanas.", "(MIAO; HOLMES, 2023; TIME, 2023; ASSOCIATED PRESS, 2023)")
body(doc, "Nesse contexto, este trabalho investiga como estruturar um artefato de apoio ao estudo que utilize IA generativa sem substituir o professor, associe questões a habilidades específicas e transforme erros em ações de revisão. A pergunta orientadora é: como projetar e verificar tecnicamente uma plataforma de estudo complementar para o Ensino Médio brasileiro que combine geração automatizada, curadoria docente, feedback formativo e acompanhamento por habilidade?")
heading(doc, "1.1 Objetivo geral", 2)
body(doc, "Projetar, implementar e verificar tecnicamente o AISTUDYTEC como artefato de apoio ao estudo independente, alinhado a princípios curriculares brasileiros e a práticas de aprendizagem baseadas em evidências, mantendo supervisão humana sobre conteúdos oficiais.")
heading(doc, "1.2 Objetivos específicos", 2)
for x in ["Modelar o ciclo de criação, geração, revisão e publicação de tópicos.", "Implementar explicações em três níveis e quizzes com habilidade, gabarito e feedback.", "Agregar desempenho por habilidade e produzir fila de revisão.", "Integrar recuperação de vídeos reais com critérios explícitos de duração e ranking.", "Verificar regras críticas por testes automatizados.", "Analisar riscos éticos, jurídicos e operacionais.", "Propor protocolo de avaliação educacional sem fabricar resultados."]:
    bullet(doc, x)
heading(doc, "1.3 Contribuição e escopo", 2)
body(doc, "A contribuição é um protótipo e sua documentação técnico-científica. Não se apresenta ensaio clínico, experimento educacional ou evidência causal de ganho de aprendizagem. A ausência desses dados não é preenchida por simulação narrativa; é tratada como limitação e como requisito para a próxima fase.")

chapter(doc, 2, "Contexto educacional e problema")
body(doc, "O PISA 2022 mostrou que 73% dos estudantes brasileiros avaliados tiveram baixo desempenho em matemática. A estabilidade dos resultados desde 2009 reforça que soluções tecnológicas isoladas não resolvem problemas estruturais; elas precisam integrar-se a currículo, trabalho docente, infraestrutura e políticas públicas.", "(INEP, 2023)")
add_table(doc, ["Indicador", "Resultado real", "Implicação para o projeto"], [
    ["Amostra brasileira no PISA 2022", "10.798 estudantes; 599 escolas", "Contextualiza, não mede usuários do AISTUDYTEC"],
    ["Baixo desempenho em matemática", "73%", "Justifica apoio diagnóstico e prática guiada"],
    ["Escolas com Internet", "92% em 2023", "Há alcance potencial, mas acesso não é uniforme"],
    ["Internet disponível aos alunos na sala", "65% das escolas conectadas", "Exige experiência tolerante a limitações"],
    ["Política de proteção de dados", "55% das escolas", "Governança precisa ser parte do produto"],
    ["Formação docente em tecnologias", "54% das escolas", "Adoção requer apoio ao professor"]
], widths=[2.0, 1.8, 2.9], caption="Tabela 1 — Indicadores educacionais e digitais usados na delimitação", source="INEP (2023) e CGI.br/NIC.br (2024).")
body(doc, "A TIC Educação 2023, baseada em 3.001 entrevistas com gestores, registrou crescimento do acesso à Internet nas escolas de Ensino Fundamental e Médio de 82% em 2020 para 92% em 2023. Contudo, apenas 30% possuíam velocidade superior a 51 Mbps e, entre escolas conectadas, 65% disponibilizavam Internet aos alunos em sala. O desenho do produto deve, portanto, evitar pressupor conectividade contínua e dispositivos abundantes.", "(CGI.br/NIC.br, 2024)")
body(doc, "A BNCC define aprendizagens essenciais por competências e habilidades. O AISTUDYTEC adota essa granularidade como inspiração operacional, mas não afirma que rótulos gerados automaticamente correspondam a códigos oficiais da BNCC. A validação curricular deve permanecer sob responsabilidade docente.", "(BRASIL, 2018)")

chapter(doc, 3, "Referencial teórico")
heading(doc, "3.1 Avaliação formativa, recuperação e feedback", 2)
body(doc, "Testes podem funcionar como oportunidades de aprendizagem, e não apenas de mensuração. Revisões da literatura mostram benefícios da prática de recuperação e do feedback em diferentes idades e conteúdos. Kubik, Gaschler e Hausman sintetizam que a recuperação ativa acompanhada de feedback constitui ferramenta robusta, embora moderadores de contexto, material e implementação devam ser considerados.", "(KUBIK; GASCHLER; HAUSMAN, 2021)")
body(doc, "No AISTUDYTEC, cada pergunta produz feedback imediato no cliente e o servidor recalcula a nota a partir do gabarito persistido. Essa escolha reduz confiança em pontuação enviada pelo navegador e aproxima a interação de uma avaliação formativa. Ainda assim, a qualidade do feedback depende da questão e da explicação geradas, razão pela qual o professor revisa o conteúdo oficial.")
heading(doc, "3.2 Prática espaçada", 2)
body(doc, "A prática distribuída apresenta evidência consistente, mas o intervalo ótimo varia conforme domínio, retenção desejada e desenho da atividade. Meta-análise sobre aprendizagem de segunda língua encontrou efeito de médio a grande e diferenças entre avaliações imediatas e tardias. O intervalo fixo de três dias do protótipo é uma decisão de produto a ser empiricamente avaliada, não uma constante universal derivada da literatura.", "(KIM; WEBB, 2022)")
heading(doc, "3.3 Sistemas tutores e personalização", 2)
body(doc, "Meta-análise recente de 30 estudos sobre sistemas tutores inteligentes reportou efeito agregado positivo, mas resultados menos conclusivos para aquisição de conhecimento, motivação, desempenho e resolução de problemas, além de variação por desenho e contexto. Essa heterogeneidade impede extrapolar eficácia para o AISTUDYTEC sem estudo próprio.", "(HUANG; XU; LIU, 2025)")
heading(doc, "3.4 IA generativa com supervisão humana", 2)
body(doc, "A orientação da UNESCO enfatiza agência humana, inclusão, diversidade e proteção de dados. O gate de publicação do AISTUDYTEC materializa parte dessa orientação: a IA propõe; o professor edita, aprova vídeos e autoriza a publicação. Temas livres do aluno são identificados como não revisados e não se tornam conteúdo oficial.", "(MIAO; HOLMES, 2023)")

chapter(doc, 4, "Trabalhos e debates relacionados")
body(doc, "O debate público registra usos de chatbots para explicações e feedback, mas também ressalta que tecnologias não substituem vínculos pedagógicos. Reportagem da Associated Press descreveu professores de matemática que combinam domínio conceitual, discussão de limitações e uso instrumental da IA. A revista Time registrou posições de pesquisadores de Harvard e Columbia sobre ciclos de entusiasmo, desigualdade e vieses de representação. Esses relatos contextualizam práticas e controvérsias; não são usados como prova causal.", "(ASSOCIATED PRESS, 2023; TIME, 2023)")
add_table(doc, ["Abordagem", "Força", "Risco", "Resposta no AISTUDYTEC"], [
    ["Chatbot direto", "Interação flexível", "Erro e ausência de curadoria", "Conteúdo oficial revisado"],
    ["Quiz convencional", "Medição simples", "Nota sem diagnóstico", "Habilidade e feedback por item"],
    ["Vídeo recomendado", "Multimodalidade", "Qualidade variável", "Filtro + aprovação por nível"],
    ["Tutoria adaptativa", "Personalização", "Opacidade e dados", "Regras simples e auditáveis"]
], widths=[1.4, 1.7, 1.8, 2.0], caption="Tabela 2 — Posicionamento conceitual do artefato", source="Elaboração própria com base na arquitetura implementada e na literatura citada.")

chapter(doc, 5, "Metodologia")
body(doc, "A pesquisa é aplicada, qualitativa quanto à construção do artefato e quantitativa apenas nas verificações técnicas descritivas. O percurso inspira-se em Design Science Research: identificação do problema, definição de objetivos, projeto e desenvolvimento, demonstração técnica, avaliação limitada e comunicação. Não houve coleta com participantes humanos.")
heading(doc, "5.1 Fontes e procedimentos", 2)
for x in ["Análise documental da BNCC, PISA 2022, TIC Educação 2023, UNESCO e legislação brasileira.", "Revisão narrativa orientada por artigos revisados por pares sobre recuperação, feedback, espaçamento e tutoria inteligente.", "Inspeção direta do repositório e contagem reproduzível de rotas, tabelas, componentes e linhas.", "Execução de testes automatizados do backend e compilação do frontend.", "Análise de riscos técnicos, pedagógicos e de privacidade."]:
    bullet(doc, x)
heading(doc, "5.2 Critérios de validade", 2)
body(doc, "Validade técnica foi limitada aos comportamentos cobertos pelos testes. Validade de conteúdo exige avaliação por professores das disciplinas. Validade educacional requer estudo com desenho pré-registrado, instrumentos adequados e comparação temporal ou entre condições. Validade externa não pode ser inferida nesta etapa.")
heading(doc, "5.3 Reprodutibilidade", 2)
body(doc, "Os testes podem ser executados por `python -m pytest backend/tests -q`; o frontend por `npm.cmd run build`. O documento registra a data da medição e distingue contagens do repositório de dados educacionais externos.")

chapter(doc, 6, "Especificação do AISTUDYTEC")
body(doc, "O AISTUDYTEC atende professor e aluno. O professor cria turma e tópico, aciona geração, revisa três explicações, questões e vídeos e publica após critérios mínimos. O aluno acessa tópicos publicados ou solicita tema livre, escolhe nível, estuda, responde ao quiz e consulta domínio e revisões.")
heading(doc, "6.1 Regras de negócio", 2)
for x in ["Domínio por habilidade: percentual acumulado maior ou igual a 70%.", "Revisão: uma pendência por aluno e habilidade, vencendo três dias após desempenho abaixo do limiar.", "Geração: entre oito e dez questões; implementação solicita nove.", "Publicação: três explicações, cinco questões válidas e um vídeo aprovado por nível.", "Tema livre: não revisado, vinculado ao aluno e não oficial."]:
    bullet(doc, x)
heading(doc, "6.2 Modelo de dados", 2)
add_table(doc, ["Tabela", "Responsabilidade"], [
    ["users / classes", "Atores e vínculo com turma"], ["topics", "Estado e autoria do tópico"],
    ["topic_explanations", "Conteúdo por nível"], ["topic_learning_paths", "Trilhas imersivas estruturadas"], ["quiz_questions", "Questões, habilidades e gabarito"],
    ["topic_videos", "Candidatos e aprovação"], ["history / quiz_attempt_answers", "Tentativas e respostas"],
    ["skill_mastery", "Agregação por aluno e habilidade"], ["review_queue", "Revisões pendentes"]
], widths=[2.1, 4.6], caption="Tabela 3 — Entidades persistidas", source="Inspeção de backend/app.py em 21 jul. 2026.")

chapter(doc, 7, "Arquitetura e implementação")
body(doc, "A solução utiliza React no frontend e Flask no backend, com SQLite via módulo padrão `sqlite3`. Não há ORM, roteador frontend ou gerenciador externo de estado. Gemini e YouTube são chamados no servidor, reduzindo exposição das chaves no navegador.")
add_table(doc, ["Camada", "Tecnologia", "Função"], [
    ["Interface", "React 18 + lucide-react", "Fluxos de professor e aluno"],
    ["API", "Flask + flask-cors", "Regras e contratos HTTP"],
    ["Persistência", "SQLite", "Dados do protótipo"],
    ["Geração", "Gemini", "Explicações e quiz estruturado"],
    ["Vídeo", "YouTube Data API v3", "Busca relevante, filtro de 3–20 min e ranking"]
], widths=[1.2, 2.0, 3.5], caption="Tabela 4 — Arquitetura lógica", source="Inspeção do código e arquivos de dependências.")
heading(doc, "7.1 Integridade da geração", 2)
body(doc, "O cliente Gemini solicita JSON e valida presença das três explicações, quantidade de oito a dez questões, campos obrigatórios, opções, resposta correta e dificuldade. Respostas malformadas produzem erro controlado. Essa validação é estrutural; não comprova correção factual ou adequação curricular.")
body(doc, "A baseline atual também exige três trilhas imersivas completas. Cada nível contém gancho, objetivos, ideias essenciais, conexão real, investigação com comparação de fontes, missão antes/durante/depois do vídeo, desafio prático, reflexão e discussão. O diário permanece local ao dispositivo e não integra avaliação docente.")
heading(doc, "7.2 Recuperação de vídeos", 2)
body(doc, "A busca combina tema, nível e referência ao Ensino Médio. Durações ISO 8601 são convertidas para segundos e conteúdos fora de três a vinte minutos são descartados. O ranking exige correspondência entre termos relevantes do tema e o título do vídeo, preserva relevância da API e usa popularidade apenas como sinal secundário. Resultado ausente é preferível a conteúdo fora do assunto; temas oficiais ainda dependem de aprovação docente.")
heading(doc, "7.3 Segurança", 2)
body(doc, "A baseline implementa sessão HTTP-only, papéis, autorização por recurso, defesa de origem, CORS por allowlist, rate limit, headers e auditoria minimizada. Recuperação e revogação central de sessão, CSRF definitivo, migrações versionadas e backup/restauração ainda são requisitos para exposição pública.")

chapter(doc, 8, "Verificação e resultados técnicos")
body(doc, "Em 22 de julho de 2026, a inspeção reproduzível identificou 23 rotas Flask, 12 tabelas incluindo auditoria, 13 componentes React, 40 testes backend e 15 testes frontend. A contagem é descritiva e não representa eficácia, segurança integral ou qualidade educacional.")
add_table(doc, ["Evidência", "Resultado", "Interpretação válida"], [
    ["Pytest", "40 casos backend aprovados", "Comportamentos cobertos não regrediram"],
    ["Frontend", "15 testes e build aprovados", "Componentes e derivação cobertos compilam"],
    ["Rotas", "23", "Superfície HTTP implementada"],
    ["Tabelas", "12", "Modelo persistente presente"],
    ["Componentes React", "13", "Separação básica por responsabilidade"]
], widths=[1.5, 2.0, 3.2], caption="Tabela 5 — Evidências técnicas observadas", source="Execução local e inspeção do repositório em 21 jul. 2026.")
heading(doc, "8.1 Cobertura funcional observada", 2)
body(doc, "Os 40 testes backend cobrem regras de turma, publicação, domínio, revisão, schema Gemini, parsers, sessão, autorização por papel e recurso, defesa de origem, rate limit, validação negativa e auditoria. Os 15 testes frontend cobrem primitives, progressão e modelo de intervenção. Smoke tests reais confirmaram HTTP 200 no Gemini e YouTube; uma geração real persistiu três trilhas e nove questões.")
heading(doc, "8.2 O que os resultados não demonstram", 2)
body(doc, "A aprovação dos testes não demonstra ausência de vulnerabilidades, desempenho sob carga, confiabilidade das APIs reais, acessibilidade, usabilidade, qualidade de conteúdo ou efeito educacional. O build não substitui teste de jornada no navegador. Nenhum resultado de aprendizagem foi coletado.")

chapter(doc, 9, "Discussão")
body(doc, "O principal mérito do desenho é converter geração de conteúdo em processo editorial docente. Essa arquitetura responde à recomendação de supervisão humana e reduz a probabilidade de conteúdo oficial ser publicado automaticamente. O sistema também preserva o gabarito no servidor, evitando confiar na nota calculada pelo cliente.")
body(doc, "A granularidade por habilidade possibilita diagnóstico mais acionável do que uma nota global, mas cria desafio semântico: habilidades produzidas por linguagem natural podem variar em redação e fragmentar agregações equivalentes. Uma evolução necessária é vocabulário controlado, mapeamento validado por professor e versionamento das habilidades.")
body(doc, "O intervalo fixo de três dias é transparente e auditável, porém pedagogicamente simplificado. A literatura sustenta espaçamento, não um intervalo universal. A agenda científica deve comparar retenção e adesão sob diferentes regras sem alterar parâmetros durante o piloto sem protocolo.")
body(doc, "A TIC Educação 2023 recomenda cautela com pressupostos de infraestrutura. A experiência deve funcionar em telas pequenas, reduzir chamadas redundantes, tolerar rede instável e permitir ao professor antecipar conteúdo quando possível. A ausência atual de modo offline é relevante para equidade.")

chapter(doc, 10, "Ética, privacidade e riscos")
body(doc, "A LGPD protege liberdade, privacidade e desenvolvimento da personalidade e exige atenção especial ao tratamento de dados de crianças e adolescentes. A ANPD afirma que o melhor interesse deve prevalecer em qualquer hipótese. O projeto armazena identidade, vínculo escolar e desempenho, dados que podem produzir inferências sensíveis sobre trajetória acadêmica.", "(BRASIL, 2018b; ANPD, 2023)")
add_table(doc, ["Risco", "Consequência", "Mitigação prioritária"], [
    ["Alucinação da IA", "Conteúdo incorreto", "Curadoria, validação e rastreabilidade"],
    ["Acesso indevido", "Exposição de dados", "Autorização por papel e testes"],
    ["Viés", "Tratamento desigual", "Revisão humana e avaliação por grupos"],
    ["Dependência de API", "Indisponibilidade/custo", "Timeout, erro claro e limites"],
    ["Conectividade", "Exclusão", "Design leve e estratégia offline"],
    ["Automação excessiva", "Redução da agência", "Professor como publicador final"]
], widths=[1.4, 2.0, 3.3], caption="Tabela 6 — Registro inicial de riscos", source="Elaboração própria com base em UNESCO (2023), LGPD e arquitetura.")
body(doc, "Antes de qualquer coleta, recomenda-se inventário de dados, finalidade e base legal; política de retenção; controle de acesso; consentimento quando aplicável; canal para titulares; avaliação de fornecedores; plano de incidentes; e submissão ao comitê de ética quando a pesquisa envolver participantes.")

chapter(doc, 11, "Protocolo de avaliação futura")
body(doc, "Propõe-se estudo piloto de métodos mistos, sem afirmar execução. A unidade de análise pode ser o estudante, mas a alocação deve considerar contaminação por turma. Uma alternativa é desenho quase-experimental com pré-teste e pós-teste, grupo de comparação equivalente e medidas de retenção tardia. Se viável ética e logisticamente, randomização por turma reduz troca de intervenção entre colegas.")
heading(doc, "11.1 Hipóteses e desfechos", 2)
for x in ["H1: estudantes com acesso ao fluxo completo apresentam maior ganho em teste alinhado às habilidades do que comparação sob prática usual.", "H2: feedback imediato reduz repetição de erros na mesma habilidade.", "H3: conclusão de revisões associa-se a maior retenção tardia.", "Desfecho primário: diferença de ganho padronizado em instrumento validado por especialistas.", "Secundários: adesão, tempo, revisões concluídas, usabilidade e percepção docente."]:
    bullet(doc, x)
heading(doc, "11.2 Integridade metodológica", 2)
body(doc, "O tamanho amostral deve ser calculado a priori com efeito mínimo relevante, poder e estrutura de agrupamento definidos. Instrumentos devem passar por validação de conteúdo. O plano de análise deve ser pré-registrado, incluindo tratamento de perdas e dados ausentes. Métricas de plataforma não podem substituir aprendizagem.")
heading(doc, "11.3 Aspectos éticos", 2)
body(doc, "Participação voluntária, assentimento do estudante e consentimento responsável quando aplicável; ausência de prejuízo escolar por recusa; minimização e pseudonimização; acesso restrito; retenção limitada; possibilidade de retirada; e comunicação de limitações da IA são requisitos mínimos.")

chapter(doc, 12, "Gestão Scrum e evolução MCP")
body(doc, "A evolução foi organizada por backlog Scrum, com objetivo de produto, histórias, critérios de aceitação, estimativas relativas e Definition of Done. O Scrum é empregado como sistema de inspeção e adaptação, não como evidência científica. Os itens prioritários são validação com provedores reais, decisão de identidade/autorização e definição da fonte do ano escolar do aluno.")
body(doc, "A integração via Model Context Protocol foi documentada como estratégia futura. O primeiro incremento deve ser local e somente leitura, expondo visão de turma ou revisão de tópico. Ferramentas mutáveis exigem identidade, autorização, confirmação e auditoria. O MCP não deve acessar SQLite diretamente; deve reutilizar regras do backend.")
add_table(doc, ["Prioridade", "Incremento", "Gate"], [
    ["Concluído", "Validar Gemini e YouTube reais", "Smoke tests HTTP 200"],
    ["P0", "Identidade e autorização", "Decisão arquitetural"],
    ["P0", "Ano escolar do aluno", "Decisão de produto"],
    ["P1", "UX de salvamento e carregamento", "Testes de jornada"],
    ["P2", "MCP somente leitura", "Autorização aprovada"]
], widths=[0.8, 3.2, 2.7], caption="Tabela 7 — Roadmap resumido", source="Product Backlog do repositório, 21 jul. 2026.")

chapter(doc, 13, "Limitações")
for x in ["Os provedores reais foram validados por smoke tests, mas não houve teste prolongado de cota, indisponibilidade ou custo.", "Não houve participantes humanos nem dados próprios de aprendizagem ou usabilidade.", "A segurança implementada ainda não é suficiente para implantação pública com menores.", "SQLite, Flask de desenvolvimento e geração síncrona não foram submetidos a carga ou alta concorrência.", "A revisão bibliográfica é narrativa, não revisão sistemática com protocolo PRISMA.", "As tags de habilidade não são códigos oficiais da BNCC e exigem validação docente.", "O intervalo fixo de revisão não foi comparado empiricamente.", "Acessibilidade e experiência em conectividade instável não foram formalmente auditadas."]:
    bullet(doc, x)
body(doc, "Essas limitações restringem a conclusão ao seguinte: existe um artefato funcional com decisões pedagógicas explícitas e evidência técnica inicial. Qualquer afirmação sobre eficácia, escala ou segurança em produção seria prematura.")

chapter(doc, 14, "Conclusão")
body(doc, "O AISTUDYTEC foi concebido para um problema real: apoiar estudo complementar sem deslocar o professor. Sua arquitetura conecta geração em níveis, questões por habilidade, feedback imediato, vídeos curados, domínio acumulado e revisão. O gate docente e a separação entre conteúdo oficial e tema livre são decisões coerentes com governança humana da IA.")
body(doc, "A avaliação disponível é técnica. Os 40 testes backend, 15 testes frontend, build bem-sucedido e smoke tests reais dos provedores sustentam confiança limitada nos comportamentos examinados. Não sustentam eficácia educacional. A maturidade para piloto depende de segurança para produção, privacidade, acessibilidade, validação docente e protocolo ético.")
body(doc, "Como contribuição acadêmico-tecnológica, o trabalho oferece um artefato auditável, documentação de requisitos e riscos e um plano de avaliação que evita confundir funcionamento com aprendizagem. O próximo avanço cientificamente defensável é executar um piloto controlado, com instrumentos válidos, dados reais e análise previamente definida.")

# Referências
doc.add_page_break(); heading(doc, "REFERÊNCIAS", 1)
refs = [
"ANPD — AUTORIDADE NACIONAL DE PROTEÇÃO DE DADOS. ANPD divulga enunciado sobre o tratamento de dados pessoais de crianças e adolescentes. Brasília, 2023. Disponível em: https://www.gov.br/anpd/pt-br/assuntos/noticias/anpd-divulga-enunciado-sobre-o-tratamento-de-dados-pessoais-de-criancas-e-adolescentes. Acesso em: 21 jul. 2026.",
"ASSOCIATED PRESS. Chatbots might disrupt math and computer science classes. Some teachers see upsides. 24 out. 2023. Disponível em: https://apnews.com/article/chatgpt-math-computer-science-3fc4b72d69d34627ba3f2fa74491ea21. Acesso em: 21 jul. 2026.",
"BRASIL. Ministério da Educação. Base Nacional Comum Curricular: Ensino Médio. Brasília: MEC, 2018. Disponível em: https://www.gov.br/mec/pt-br/cne/bncc_ensino_medio.pdf. Acesso em: 21 jul. 2026.",
"BRASIL. Lei nº 13.709, de 14 de agosto de 2018. Lei Geral de Proteção de Dados Pessoais. Brasília, 2018b. Disponível em: https://www.planalto.gov.br/ccivil_03/_ato2015-2018/2018/lei/l13709.htm. Acesso em: 21 jul. 2026.",
"CGI.br/NIC.br. Pesquisa sobre o uso das tecnologias de informação e comunicação nas escolas brasileiras: TIC Educação 2023. São Paulo: Comitê Gestor da Internet no Brasil, 2024. Disponível em: https://cetic.br/media/docs/publicacoes/2/20241119194257/tic_educacao_2023_livro_completo.pdf. Acesso em: 21 jul. 2026.",
"HUANG, Xiaoli; XU, Wei; LIU, Ruijia. Effects of Intelligent Tutoring Systems on Educational Outcomes: a meta-analysis. International Journal of Distance Education Technologies, v. 23, n. 1, 2025. DOI: 10.4018/IJDET.368420.",
"INEP — INSTITUTO NACIONAL DE ESTUDOS E PESQUISAS EDUCACIONAIS ANÍSIO TEIXEIRA. Pisa 2022: resultados do Brasil. Brasília, 2023. Disponível em: https://www.gov.br/inep/pt-br/areas-de-atuacao/avaliacao-e-exames-educacionais/pisa/resultados/2022. Acesso em: 21 jul. 2026.",
"KIM, Su Kyung; WEBB, Stuart. The Effects of Spaced Practice on Second Language Learning: A Meta-Analysis. Language Learning, v. 72, n. 1, 2022. DOI: 10.1111/lang.12479.",
"KUBIK, Veit; GASCHLER, Robert; HAUSMAN, Hannah. Enhancing Student Learning in Research and Educational Practice: The Power of Retrieval Practice and Feedback. Psychology Learning & Teaching, v. 20, n. 1, 2021. DOI: 10.1177/1475725720976462.",
"MIAO, Fengchun; HOLMES, Wayne. Guidance for generative AI in education and research. Paris: UNESCO, 2023. Disponível em: https://unesdoc.unesco.org/ark:/48223/pf0000386693. Acesso em: 21 jul. 2026.",
"OECD. PISA 2022 Results (Volume I and II) — Country Note: Brazil. Paris: OECD Publishing, 2023. Disponível em: https://www.oecd.org/en/publications/pisa-2022-results-volume-i-and-ii-country-notes_ed6fbcc5-en/brazil_61690648-en.html. Acesso em: 21 jul. 2026.",
"TIME. The Creative Ways Teachers Are Using ChatGPT in the Classroom. 8 ago. 2023. Disponível em: https://time.com/6300950/ai-schools-chatgpt-teachers/. Acesso em: 21 jul. 2026."
]
for ref in refs:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8); p.paragraph_format.line_spacing = 1.0
    font(p.add_run(ref), 10.5)

# Apêndices
doc.add_page_break(); heading(doc, "APÊNDICE A — CONTRATOS FUNCIONAIS RESUMIDOS", 1)
add_table(doc, ["Grupo", "Operações"], [
    ["Acesso", "login, cadastro, entrada em turma"],
    ["Professor", "turmas, alunos, tópicos, geração, edição e publicação"],
    ["Aluno", "tópicos publicados, tema livre, estudo e progresso"],
    ["Avaliação", "tentativa, respostas, domínio, revisão e dashboard"]
], widths=[1.5, 5.2], caption="Tabela A1 — Superfície funcional", source="Inspeção das 23 rotas Flask.")
heading(doc, "Critérios técnicos de publicação", 2)
for x in ["Explicações `simple`, `technical` e `advanced` não vazias.", "Pelo menos cinco questões com habilidade e resposta correta.", "Pelo menos um vídeo aprovado em cada nível.", "Erro 400 com lista de pendências quando incompleto."]:
    bullet(doc, x)

doc.add_page_break(); heading(doc, "APÊNDICE B — MATRIZ DE RASTREABILIDADE", 1)
add_table(doc, ["Requisito", "Implementação", "Evidência atual"], [
    ["Curadoria docente", "Gate de publicação", "Teste de rejeição/sucesso"],
    ["Feedback imediato", "Componente Quiz", "Inspeção de código; falta teste E2E"],
    ["Domínio por habilidade", "skill_mastery", "Teste de upsert"],
    ["Revisão espaçada", "review_queue", "Teste de deduplicação"],
    ["JSON estruturado", "validador Gemini", "Testes válidos/inválidos"],
    ["Vídeos 3–20 min e tema no título", "parser, filtro e ranking", "Testes do parser e smoke test real"],
    ["Trilha imersiva", "learningPaths + StudyView", "Geração real, persistência e build"]
], widths=[1.6, 2.1, 3.0], caption="Tabela B1 — Requisitos e evidências", source="Testes e código do repositório em 22 jul. 2026.")

doc.add_page_break(); heading(doc, "APÊNDICE C — CHECKLIST PARA PILOTO", 1)
for x in ["Aprovação ética e institucional quando aplicável.", "Identidade, autorização e política de senhas.", "Inventário LGPD, base legal, retenção e direitos dos titulares.", "Validação real de Gemini e YouTube sem vazamento de chaves.", "Revisão curricular por professores das disciplinas.", "Auditoria de acessibilidade e conectividade limitada.", "Instrumentos de pré, pós e retenção validados.", "Cálculo amostral, pré-registro e plano de análise.", "Canal de suporte e resposta a incidentes.", "Critérios de interrupção diante de conteúdo danoso."]:
    bullet(doc, x)

doc.core_properties.title = "AISTUDYTEC — Trabalho técnico-científico"
doc.core_properties.subject = "IA generativa, curadoria docente e estudo por habilidades"
doc.core_properties.author = "Projeto AISTUDYTEC"
doc.core_properties.keywords = "AISTUDYTEC, educação, inteligência artificial, Ensino Médio"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
